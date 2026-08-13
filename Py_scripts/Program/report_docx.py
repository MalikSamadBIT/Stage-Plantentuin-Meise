import io
import math
from datetime import datetime

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt, RGBColor
from PIL import Image as PILImage

import gap_analysis

# Same status vocabulary/colors as report_pdf.py, for visual parity between
# the two export formats.
STATUS_COLORS = {
    "no_sequences": RGBColor(0xB9, 0x1C, 0x1C),
    "partial": RGBColor(0xB4, 0x53, 0x09),
    "complete": RGBColor(0x15, 0x80, 0x3D),
    "no_data": RGBColor(0x6B, 0x72, 0x80),
}
MUTED_COLOR = RGBColor(0x55, 0x55, 0x55)


def _formatted_paragraph(
    doc, text, size=10.5, bold=False, italic=False, alignment=None,
    color=None, space_before=0, space_after=6
):
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)

    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color

    return p


def _add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p_border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    p_border.append(bottom)
    p._p.get_or_add_pPr().append(p_border)


def _add_page_number_footer(doc):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run("Page ")
    run.font.name = "Times New Roman"
    run.font.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED_COLOR

    field_run = p.add_run()
    field_run.font.name = "Times New Roman"
    field_run.font.italic = True
    field_run.font.size = Pt(8)
    field_run.font.color.rgb = MUTED_COLOR

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    field_run._r.append(fld_begin)
    field_run._r.append(instr)
    field_run._r.append(fld_end)


def _set_cell_text(cell, text, bold=False, italic=False, color=None, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def _set_column_widths(table, widths):
    # python-docx needs the width set on every cell in a column, not just
    # table.columns[i].width, or Word ignores it
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def _gap_table_col_widths(num_markers, available_width):
    species_w = available_width * 0.26
    obs_w = available_width * 0.13
    status_w = available_width * 0.17
    marker_total = max(available_width - species_w - obs_w - status_w, Cm(1))
    marker_w = marker_total / num_markers if num_markers else 0
    return [species_w, obs_w] + [marker_w] * num_markers + [status_w]


def _build_gap_table(doc, rows, target_markers, available_width):
    table = doc.add_table(rows=1, cols=2 + len(target_markers) + 1)
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    headers = ["Species", "Observations"] + target_markers + ["Status"]
    for i, header in enumerate(headers):
        _set_cell_text(header_cells[i], header, bold=True)

    for row in rows:
        cells = table.add_row().cells
        _set_cell_text(cells[0], row["species"])

        obs = "-" if row["observations"] is None else f"{row['observations']:,}"
        _set_cell_text(cells[1], obs)

        for i, marker in enumerate(target_markers):
            value = row[marker]
            _set_cell_text(cells[2 + i], "-" if value is None else value)

        status_col = 2 + len(target_markers)
        color = STATUS_COLORS.get(row["status"])
        _set_cell_text(
            cells[status_col], row["status_label"],
            bold=row["status"] == "no_sequences",
            italic=row["status"] == "partial",
            color=color
        )

    _set_column_widths(table, _gap_table_col_widths(len(target_markers), available_width))
    return table


def _build_marker_coverage_table(doc, coverage_stats, available_width):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    _set_cell_text(table.rows[0].cells[0], "Marker", bold=True)
    _set_cell_text(table.rows[0].cells[1], "Species covered", bold=True)
    _set_cell_text(table.rows[0].cells[2], "Coverage", bold=True)

    for marker, covered, assessed, percentage in coverage_stats:
        cells = table.add_row().cells
        _set_cell_text(cells[0], marker)
        _set_cell_text(cells[1], f"{covered} / {assessed}")
        _set_cell_text(cells[2], f"{percentage:.0f}%")

    _set_column_widths(
        table, [available_width * 0.4, available_width * 0.3, available_width * 0.3])
    return table


def _build_score_table(doc, scores, available_width):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    _set_cell_text(table.rows[0].cells[0], "Metric", bold=True)
    _set_cell_text(table.rows[0].cells[1], "Value", bold=True)

    for label, value in scores.items():
        if "% of max possible" in label:
            formatted = f"{value:.2f}%"
        elif isinstance(value, float):
            formatted = f"{value:.4f}"
        else:
            formatted = str(value)
        cells = table.add_row().cells
        _set_cell_text(cells[0], label)
        _set_cell_text(cells[1], formatted)

    _set_column_widths(table, [available_width * 0.7, available_width * 0.3])
    return table


def _build_synonyms_table(doc, synonym_results, available_width):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    _set_cell_text(table.rows[0].cells[0], "Species", bold=True)
    _set_cell_text(table.rows[0].cells[1], "Synonyms", bold=True)

    for species, synonyms in synonym_results.items():
        parts = [f"{name} ({lang})" for lang, name in synonyms.items() if name]
        cells = table.add_row().cells
        _set_cell_text(cells[0], species)
        _set_cell_text(cells[1], ", ".join(parts) if parts else "-")

    _set_column_widths(table, [available_width * 0.28, available_width * 0.72])
    return table


def _add_split_image(doc, image_bytes, available_width, target_dpi=150):
    # Same problem/fix as report_pdf.py's _split_wide_image: a long MSA
    # alignment scaled down to page width becomes unreadable, so crop it
    # into width-limited strips (full height) at a consistent DPI instead.
    img = PILImage.open(io.BytesIO(image_bytes))
    native_width, native_height = img.size

    max_px_per_strip = available_width.inches * target_dpi

    if native_width <= max_px_per_strip:
        doc.add_picture(io.BytesIO(image_bytes), width=available_width)
        return

    num_strips = math.ceil(native_width / max_px_per_strip)
    strip_width = math.ceil(native_width / num_strips)

    for i in range(num_strips):
        left = i * strip_width
        right = min(left + strip_width, native_width)
        strip = img.crop((left, 0, right, native_height))

        buf = io.BytesIO()
        strip.save(buf, format="PNG")
        buf.seek(0)

        doc.add_picture(buf, width=available_width)


def build_report_docx(
    path, rows, target_markers, meta, report_items, status_chart_bytes=None
):
    """
    Same inputs/shape as report_pdf.build_report_pdf - see that docstring.
    """
    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(10.5)

    section = doc.sections[0]
    if len(target_markers) > 4:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = (
            section.page_height, section.page_width
        )
    section.left_margin = section.right_margin = Cm(2.2)
    section.top_margin = section.bottom_margin = Cm(2)

    # python-docx's Length arithmetic returns a plain int, not a Length
    # subclass - re-wrap so .inches etc. stay available downstream
    available_width = Emu(
        section.page_width - section.left_margin - section.right_margin)

    _formatted_paragraph(
        doc, "Gap analysis report", size=20, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2
    )
    _formatted_paragraph(
        doc,
        f"Flora Fetch - generated {datetime.now().strftime('%d %B %Y')}",
        size=10, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        color=MUTED_COLOR, space_after=2
    )
    _formatted_paragraph(
        doc,
        f"Site: {meta['site']} - period {meta['start_date']} to "
        f"{meta['end_date']} - markers: {', '.join(target_markers)}",
        size=10, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        color=MUTED_COLOR, space_after=4
    )
    _add_horizontal_rule(doc)

    counts = {"no_sequences": 0, "partial": 0, "complete": 0, "no_data": 0}
    for row in rows:
        counts[row["status"]] += 1

    no_data_note = (
        f" ({counts['no_data']} could not be checked against the "
        f"observation site)." if counts["no_data"] else "."
    )
    abstract = (
        f"Of the {len(rows)} species surveyed, {counts['no_sequences']} lack "
        f"any sequence data on file, {counts['partial']} show partial marker "
        f"coverage, and {counts['complete']} are fully represented in the "
        f"local reference database{no_data_note} Within each coverage "
        "category, species are ranked by field-observation count, so the "
        "highest-value near-term sequencing targets appear first."
    )
    _formatted_paragraph(
        doc, abstract, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_before=8, space_after=10
    )

    _formatted_paragraph(doc, "1. Gap analysis", size=13, bold=True,
                          space_before=10, space_after=2)
    _formatted_paragraph(
        doc,
        "Table 1. Sequence coverage per species, ranked by observation "
        "count within status.",
        size=9, italic=True, color=MUTED_COLOR, space_after=4
    )
    _build_gap_table(doc, rows, target_markers, available_width)

    coverage_stats = gap_analysis.build_marker_coverage_stats(rows, target_markers)
    _formatted_paragraph(
        doc,
        "Table 2. Marker coverage across surveyed species (of those "
        "successfully checked against the observation site).",
        size=9, italic=True, color=MUTED_COLOR, space_after=4
    )
    _build_marker_coverage_table(doc, coverage_stats, available_width)

    figure_num = 1
    if status_chart_bytes:
        doc.add_picture(io.BytesIO(status_chart_bytes), width=available_width)
        _formatted_paragraph(
            doc, f"Figure {figure_num}. Gap analysis status summary.",
            size=9, italic=True, color=MUTED_COLOR, space_after=4
        )
        figure_num += 1

    section_num = 2
    table_num = 3

    for item in report_items:
        if item["type"] == "retrieval_chart":
            _formatted_paragraph(doc, f"{section_num}. Retrieval rate",
                                  size=13, bold=True, space_before=14, space_after=2)
            doc.add_picture(io.BytesIO(item["image_bytes"]), width=available_width)
            _formatted_paragraph(
                doc, f"Figure {figure_num}. {item['title']}.",
                size=9, italic=True, color=MUTED_COLOR, space_after=4
            )
            figure_num += 1
            section_num += 1

        elif item["type"] == "msa":
            _formatted_paragraph(
                doc, f"{section_num}. Multiple sequence alignment",
                size=13, bold=True, space_before=14, space_after=2
            )
            _add_split_image(doc, item["image_bytes"], available_width)
            _formatted_paragraph(
                doc, f"Figure {figure_num}. {item['title']}.",
                size=9, italic=True, color=MUTED_COLOR, space_after=4
            )
            figure_num += 1

            if item.get("scores"):
                _formatted_paragraph(
                    doc, f"Table {table_num}. Alignment quality scores.",
                    size=9, italic=True, color=MUTED_COLOR, space_after=4
                )
                _build_score_table(doc, item["scores"], available_width)
                table_num += 1

            section_num += 1

        elif item["type"] == "barcode_gap":
            a = item["assessment"]
            _formatted_paragraph(
                doc, f"{section_num}. Barcode gap - {a['species']}",
                size=13, bold=True, space_before=14, space_after=2
            )

            verdict_label = gap_analysis.BARCODE_GAP_VERDICTS[a["verdict"]]
            method = a.get("distance_method", gap_analysis.DEFAULT_DISTANCE_METHOD)
            if a["verdict"] == "insufficient_data":
                summary = f"Verdict: {verdict_label} - {a['reason']}"
            else:
                summary = (
                    f"Verdict: {verdict_label} ({method}). Max intraspecific "
                    f"distance {a['max_intraspecific']:.4f}, min interspecific "
                    f"distance {a['min_interspecific']:.4f} (nearest neighbor: "
                    f"{a['nearest_neighbor']}), gap {a['gap']:+.4f}."
                )
            _formatted_paragraph(doc, summary, space_before=4, space_after=8)

            if item.get("image_bytes"):
                doc.add_picture(io.BytesIO(item["image_bytes"]), width=available_width)
                _formatted_paragraph(
                    doc, f"Figure {figure_num}. {item['title']} - neighbor-joining "
                    "guide tree (p-distance, no bootstrap support - not a "
                    "rigorous phylogeny).",
                    size=9, italic=True, color=MUTED_COLOR, space_after=4
                )
                figure_num += 1

            section_num += 1

        elif item["type"] == "synonyms":
            _formatted_paragraph(doc, f"{section_num}. Synonyms", size=13,
                                  bold=True, space_before=14, space_after=2)
            _formatted_paragraph(
                doc,
                f"Table {table_num}. Species and associated common names "
                "by language.",
                size=9, italic=True, color=MUTED_COLOR, space_after=4
            )
            _build_synonyms_table(doc, item["data"], available_width)
            table_num += 1
            section_num += 1

    _add_page_number_footer(doc)

    doc.save(path)
